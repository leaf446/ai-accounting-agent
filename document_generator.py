# document/document_generator.py
"""
전문 DOCX 보고서 생성 엔진 - 완성 버전
회계법인 수준의 고급 문서 템플릿
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
import os
import pandas as pd
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
import io
import base64
from visualization_engine import FinancialVisualizationEngine # Added import

class ProfessionalReportGenerator:
    """전문 회계 보고서 생성기"""
    
    def __init__(self, save_directory: str = "analysis_results/documents"):
        self.save_directory = save_directory
        self.document_templates = self._initialize_templates()
        self.color_scheme = self._initialize_colors()
        self.viz_engine = FinancialVisualizationEngine() # Added instantiation
        
        # 디렉토리 생성
        os.makedirs(save_directory, exist_ok=True)

    def _initialize_templates(self) -> Dict[str, Dict]:
        """문서 템플릿 초기화"""
        return {
            "executive_summary": {
                "title": "경영진 요약 보고서",
                "subtitle": "Executive Summary Report",
                "sections": [
                    "요약", "재무현황", "위험분석", "권고사항"
                ]
            },
            "comprehensive_audit": {
                "title": "종합 감사 분석 보고서", 
                "subtitle": "Comprehensive Audit Analysis Report",
                "sections": [
                    "감사개요", "재무제표분석", "내부통제검토", 
                    "부정위험평가", "감사의견", "개선권고"
                ]
            },
            "financial_analysis": {
                "title": "재무 분석 보고서",
                "subtitle": "Financial Analysis Report", 
                "sections": [
                    "분석개요", "재무성과", "재무상태", "현금흐름", 
                    "비율분석", "벤치마킹", "결론"
                ]
            }
        }

    def _initialize_colors(self) -> Dict[str, RGBColor]:
        """색상 스키마 초기화"""
        return {
            "primary": RGBColor(46, 134, 171),      # 전문적 파랑
            "secondary": RGBColor(70, 130, 180),    # 연한 파랑  
            "accent": RGBColor(255, 140, 0),        # 주황 (강조)
            "success": RGBColor(76, 175, 80),       # 초록 (양호)
            "warning": RGBColor(255, 152, 0),       # 주황 (주의)
            "danger": RGBColor(244, 67, 54),        # 빨강 (위험)
            "text_primary": RGBColor(33, 33, 33),   # 진한 회색
            "text_secondary": RGBColor(117, 117, 117) # 연한 회색
        }

    def generate_comprehensive_report(self, analysis_data: Dict, company_name: str) -> str:
        """종합 보고서 생성 - 메인 함수"""
        
        print(f"📋 {company_name} 종합 보고서 생성 시작...")
        
        # 새 문서 생성
        doc = Document()
        
        # 문서 기본 설정
        self._setup_document_format(doc)
        
        # 분석 기준 연도: 실제 수집된 다년도 데이터의 최신 연도 사용
        multi_year = analysis_data.get("multi_year_data", {}) or {}
        basis_year = max(multi_year.keys()) if multi_year else str(datetime.now().year - 1)

        # 표지 생성
        self._create_cover_page(doc, company_name, "종합 재무 분석 보고서",
                               "Comprehensive Financial Analysis Report", basis_year)
        
        # 페이지 나누기
        doc.add_page_break()
        
        # 목차 생성
        self._create_table_of_contents(doc)
        doc.add_page_break()
        
        # 1. 요약 섹션
        self._create_executive_summary_section(doc, analysis_data, company_name)
        doc.add_page_break()
        
        # 2. 재무현황 섹션
        self._create_financial_status_section(doc, analysis_data, company_name)
        doc.add_page_break()
        
        # 3. 위험분석 섹션  
        self._create_risk_analysis_section(doc, analysis_data, company_name)
        doc.add_page_break()
        
        # 4. 권고사항 섹션
        self._create_recommendations_section(doc, analysis_data, company_name)
        doc.add_page_break()
        
        # 5. 부록
        self._create_appendix_section(doc, analysis_data)
        
        # 파일 저장
        filename = f"{company_name}_종합분석보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.save_directory, filename)
        doc.save(filepath)
        
        print(f"✅ 보고서 생성 완료: {filepath}")
        return filepath

    def _setup_document_format(self, doc: Document):
        """문서 기본 포맷 설정"""
        
        # 페이지 설정
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)

    def _create_cover_page(self, doc: Document, company_name: str,
                          title: str, subtitle: str, basis_year: str = None):
        """표지 페이지 생성"""
        if basis_year is None:
            basis_year = str(datetime.now().year - 1)
        
        # 로고 공간 (텍스트로 대체)
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run("🏛️ PROFESSIONAL AUDIT FIRM")
        logo_run.font.size = Pt(20)
        logo_run.font.bold = True
        logo_run.font.color.rgb = self.color_scheme["primary"]
        
        # 공백 (과다한 공백은 표지가 2페이지로 넘치는 원인이 되므로 최소화)
        for _ in range(2):
            doc.add_paragraph()

        # 메인 제목
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.name = 'Malgun Gothic'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.color_scheme["primary"]
        
        # 영문 부제목
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.name = 'Calibri'
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = self.color_scheme["secondary"]
        
        # 공백
        for _ in range(2):
            doc.add_paragraph()
        
        # 회사명
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(f"분석 대상: {company_name}")
        company_run.font.name = 'Malgun Gothic'
        company_run.font.size = Pt(20)
        company_run.font.bold = True
        company_run.font.color.rgb = self.color_scheme["accent"]
        
        # 공백
        for _ in range(2):
            doc.add_paragraph()

        # 보고서 정보 테이블
        info_table = doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 내용
        info_data = [
            ("보고서 생성일", datetime.now().strftime("%Y년 %m월 %d일")),
            ("분석 기준", f"{basis_year} 회계연도 사업보고서"),
            ("데이터 출처", "DART 전자공시시스템"),
            ("생성 시스템", "AI 기반 자동 분석 시스템")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            label_cell.text = label
            value_cell.text = value
            
            # 셀 스타일링
            for cell in [label_cell, value_cell]:
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(12)
                if cell == label_cell:
                    run.font.bold = True
        
        # 하단 면책조항
        for _ in range(2):
            doc.add_paragraph()
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run = disclaimer_para.add_run(
            "본 보고서는 AI 기반 분석 결과이며, 투자 및 경영 의사결정 시 전문가 검토가 필요합니다."
        )
        disclaimer_run.font.name = 'Malgun Gothic'
        disclaimer_run.font.size = Pt(10)
        disclaimer_run.font.color.rgb = self.color_scheme["text_secondary"]

    def _create_table_of_contents(self, doc: Document):
        """목차 생성 (정적 방식)

        Word TOC 필드는 본문 제목이 Heading 스타일이 아니면 F9로도 채워지지 않고,
        사용자가 필드 갱신을 해야 하는 부담이 있음. 섹션 구성을 생성 시점에
        이미 알고 있으므로 바로 보이는 정적 목차를 삽입한다.
        """
        # 목차 제목
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run("목 차")
        toc_run.font.name = 'Malgun Gothic'
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_run.font.color.rgb = self.color_scheme["primary"]

        doc.add_paragraph()

        toc_entries = [
            ("1. 요약", "Executive Summary"),
            ("2. 재무현황 분석", "Financial Status Analysis"),
            ("3. 위험 분석", "Risk Analysis"),
            ("4. 권고사항", "Recommendations"),
            ("5. 부록", "Appendix"),
        ]

        for korean, english in toc_entries:
            entry = doc.add_paragraph()
            entry.paragraph_format.space_after = Pt(14)
            entry.paragraph_format.left_indent = Inches(1.0)

            run = entry.add_run(korean)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.color_scheme["text_primary"]

            sub_run = entry.add_run(f"  ·  {english}")
            sub_run.font.name = 'Calibri'
            sub_run.font.size = Pt(11)
            sub_run.font.italic = True
            sub_run.font.color.rgb = self.color_scheme["text_secondary"]

    def _create_executive_summary_section(self, doc: Document, analysis_data: Dict, company_name: str):
        """경영진 요약 섹션"""
        
        # 섹션 제목
        self._add_section_heading(doc, "1. 경영진 요약", "Executive Summary")
        
        # 전반적 평가
        overall_para = doc.add_paragraph()
        overall_run = overall_para.add_run("■ 전반적 평가")
        overall_run.font.name = 'Malgun Gothic'
        overall_run.font.size = Pt(14)
        overall_run.font.bold = True
        overall_run.font.color.rgb = self.color_scheme["primary"]
        
        # 평가 내용
        ratios = analysis_data.get('ratios', {})
        fraud_ratios = analysis_data.get('fraud_ratios', {})
        financial_data = analysis_data.get('financial_data', {})
        
        # 주요 지표 추출
        roe = ratios.get('ROE', 0)
        debt_ratio = ratios.get('부채비율', 100)
        fraud_score = fraud_ratios.get('종합_부정위험점수', 0)
        revenue = financial_data.get('revenue', 0)
        net_income = financial_data.get('net_income', 0)
        
        # 종합 등급 계산
        if roe > 15 and debt_ratio < 100 and fraud_score < 30:
            grade = "A (우수)"
            grade_color = self.color_scheme["success"]
        elif roe > 10 and debt_ratio < 150 and fraud_score < 50:
            grade = "B (양호)"
            grade_color = self.color_scheme["warning"]
        else:
            grade = "C (개선필요)"
            grade_color = self.color_scheme["danger"]
        
        grade_para = doc.add_paragraph()
        grade_text = f"{company_name}의 종합 재무 등급: "
        grade_para.add_run(grade_text).font.name = 'Malgun Gothic'
        grade_run = grade_para.add_run(grade)
        grade_run.font.name = 'Malgun Gothic'
        grade_run.font.bold = True
        grade_run.font.color.rgb = grade_color
        
        # 핵심 발견사항
        doc.add_paragraph()
        findings_para = doc.add_paragraph()
        findings_run = findings_para.add_run("■ 핵심 발견사항")
        findings_run.font.name = 'Malgun Gothic'
        findings_run.font.size = Pt(14)
        findings_run.font.bold = True
        findings_run.font.color.rgb = self.color_scheme["primary"]
        
        # 주요 지표 요약 테이블
        summary_table = doc.add_table(rows=6, cols=3)
        summary_table.style = 'Table Grid'
        
        # 테이블 헤더
        header_cells = summary_table.rows[0].cells
        headers = ["구분", "수치", "평가"]
        for i, header in enumerate(headers):
            cell = header_cells[i]
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(header)
            run.font.bold = True
            run.font.name = 'Malgun Gothic'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 테이블 데이터
        table_data = [
            ("ROE", f"{roe:.1f}%", "우수" if roe > 15 else "보통" if roe > 8 else "개선필요"),
            ("부채비율", f"{debt_ratio:.1f}%", "양호" if debt_ratio < 100 else "보통" if debt_ratio < 200 else "위험"),
            ("영업이익률", f"{ratios.get('영업이익률', 0):.1f}%", "분석결과"),
            ("순이익률", f"{ratios.get('순이익률', 0):.1f}%", "분석결과"),
            ("부정위험점수", f"{fraud_score:.0f}점", "낮음" if fraud_score < 30 else "보통" if fraud_score < 60 else "높음")
        ]
        
        for i, (metric, value, assessment) in enumerate(table_data, 1):
            row = summary_table.rows[i]
            
            # 각 셀에 데이터 입력
            for j, text in enumerate([metric, value, assessment]):
                cell = row.cells[j]
                para = cell.paragraphs[0]
                para.clear()
                run = para.add_run(text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(11)
                
                if j == 2:  # 평가 컬럼
                    if "우수" in text or "양호" in text or "낮음" in text:
                        run.font.color.rgb = self.color_scheme["success"]
                    elif "위험" in text or "높음" in text:
                        run.font.color.rgb = self.color_scheme["danger"]

    def _create_financial_status_section(self, doc: Document, analysis_data: Dict, company_name: str):
        """재무현황 섹션"""
        
        self._add_section_heading(doc, "2. 재무현황 분석", "Financial Status Analysis")
        
        financial_data = analysis_data.get('financial_data', {})
        ratios = analysis_data.get('ratios', {})
        
        # 2.1 재무성과
        self._add_subsection_heading(doc, "2.1 재무성과")
        
        revenue = financial_data.get('revenue', 0)
        net_income = financial_data.get('net_income', 0)
        
        performance_text = f"""
{company_name}의 2024년 재무성과는 다음과 같습니다:

• 매출액: {self._format_currency(revenue)}
• 순이익: {self._format_currency(net_income)}
• ROE: {ratios.get('ROE', 0):.2f}%
• ROA: {ratios.get('ROA', 0):.2f}%

매출 규모와 수익성 지표를 종합적으로 검토한 결과, """
        
        if ratios.get('ROE', 0) > 12:
            performance_text += "수익성이 양호한 수준으로 평가됩니다."
        else:
            performance_text += "수익성 개선이 필요한 것으로 분석됩니다."
        
        performance_para = doc.add_paragraph()
        performance_para.add_run(performance_text).font.name = 'Malgun Gothic'
        
        # 2.2 재무상태
        self._add_subsection_heading(doc, "2.2 재무상태")
        
        total_assets = financial_data.get('total_assets', 0)
        total_liabilities = financial_data.get('total_liabilities', 0)
        total_equity = financial_data.get('total_equity', 0)
        
        position_text = f"""
재무상태표 분석 결과:

• 총자산: {self._format_currency(total_assets)}
• 부채총계: {self._format_currency(total_liabilities)}
• 자본총계: {self._format_currency(total_equity)}
• 부채비율: {ratios.get('부채비율', 0):.1f}%
• 자기자본비율: {ratios.get('자기자본비율', 0):.1f}%

재무구조의 건전성을 평가한 결과, """
        
        debt_ratio = ratios.get('부채비율', 0)
        if debt_ratio < 100:
            position_text += "안정적인 재무구조를 유지하고 있습니다."
        elif debt_ratio < 200:
            position_text += "보통 수준의 재무구조입니다."
        else:
            position_text += "재무구조 개선이 필요합니다."
        
        position_para = doc.add_paragraph()
        position_para.add_run(position_text).font.name = 'Malgun Gothic'

        # 재무비율 비교 차트 추가
        doc.add_paragraph()
        self._add_subsection_heading(doc, "2.4 주요 재무비율 시각화")
        chart_path = self.viz_engine.create_ratio_comparison_chart(ratios, company_name)
        if chart_path:
            self._add_image_to_document(doc, chart_path, "주요 재무비율 비교 차트")

        # 다년도 추세 분석 차트 추가
        doc.add_paragraph()
        self._add_subsection_heading(doc, "2.5 다년도 재무 추세")
        multi_year_data = analysis_data.get('multi_year_data', {})
        if multi_year_data:
            chart_path = self.viz_engine.create_trend_analysis_chart(multi_year_data, company_name)
            if chart_path:
                self._add_image_to_document(doc, chart_path, "매출 및 순이익 추세 차트")

    def _create_risk_analysis_section(self, doc: Document, analysis_data: Dict, company_name: str):
        """위험분석 섹션"""
        
        self._add_section_heading(doc, "3. 위험분석", "Risk Analysis")
        
        fraud_ratios = analysis_data.get('fraud_ratios', {})
        
        # 3.1 부정위험 분석
        self._add_subsection_heading(doc, "3.1 부정위험 분석")
        
        risk_score = fraud_ratios.get('종합_부정위험점수', 0)
        
        risk_text = f"""
{company_name}의 부정위험 분석 결과:

• 종합 부정위험 점수: {risk_score:.0f}점 (100점 만점)
• 현금흐름 대 순이익 비율: {fraud_ratios.get('현금흐름_대_순이익_비율', 0):.2f}
• 매출채권 대 매출 비율: {fraud_ratios.get('매출채권_대_매출_비율', 0):.1f}%

위험 평가: """
        
        if risk_score < 30:
            risk_text += "낮은 위험 (부정 위험 신호가 거의 발견되지 않음)"
        elif risk_score < 60:
            risk_text += "보통 위험 (일부 주의 필요 항목 존재)"
        else:
            risk_text += "높은 위험 (즉시 정밀 검토 필요)"
        
        risk_para = doc.add_paragraph()
        risk_para.add_run(risk_text).font.name = 'Malgun Gothic'
        
        # 위험 요소 상세 분석
        if fraud_ratios.get('순이익_양수_현금흐름_음수', False):
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("⚠️ 주요 위험 신호: 순이익은 양수이나 영업현금흐름이 음수입니다.")
            warning_run.font.name = 'Malgun Gothic'
            warning_run.font.color.rgb = self.color_scheme["danger"]
            warning_run.font.bold = True

        # 부정 위험 레이더 차트 추가
        doc.add_paragraph()
        self._add_subsection_heading(doc, "3.2 부정 위험 시각화")
        chart_path = self.viz_engine.create_fraud_risk_radar_chart(fraud_ratios, company_name)
        if chart_path:
            self._add_image_to_document(doc, chart_path, "부정 위험 레이더 차트")

    def _create_recommendations_section(self, doc: Document, analysis_data: Dict, company_name: str):
        """권고사항 섹션"""
        
        self._add_section_heading(doc, "4. 권고사항", "Recommendations")
        
        ratios = analysis_data.get('ratios', {})
        fraud_ratios = analysis_data.get('fraud_ratios', {})
        
        # 4.1 즉시 조치사항
        self._add_subsection_heading(doc, "4.1 즉시 조치사항")
        
        immediate_actions = []
        
        # 부정위험 관련
        risk_score = fraud_ratios.get('종합_부정위험점수', 0)
        if risk_score > 60:
            immediate_actions.append("부정위험 정밀 조사 실시")
        
        if fraud_ratios.get('순이익_양수_현금흐름_음수', False):
            immediate_actions.append("현금흐름 개선 방안 수립")
        
        # 재무비율 관련
        if ratios.get('부채비율', 0) > 200:
            immediate_actions.append("부채 구조조정 검토")
        
        if ratios.get('ROE', 0) < 5:
            immediate_actions.append("수익성 개선 전략 수립")
        
        if not immediate_actions:
            immediate_actions.append("현재 특별한 즉시 조치사항은 없습니다.")
        
        for i, action in enumerate(immediate_actions, 1):
            action_para = doc.add_paragraph()
            action_para.add_run(f"{i}. {action}").font.name = 'Malgun Gothic'
        
        # 4.2 중장기 개선사항
        doc.add_paragraph()
        self._add_subsection_heading(doc, "4.2 중장기 개선사항")
        
        longterm_actions = [
            "재무 모니터링 시스템 구축",
            "내부통제 시스템 강화", 
            "정기적 재무비율 분석 및 벤치마킹",
            "경영진 대상 재무 교육 프로그램 실시"
        ]
        
        for i, action in enumerate(longterm_actions, 1):
            action_para = doc.add_paragraph()
            action_para.add_run(f"{i}. {action}").font.name = 'Malgun Gothic'

    def _create_appendix_section(self, doc: Document, analysis_data: Dict):
        """부록 섹션"""
        
        self._add_section_heading(doc, "부록 A. 상세 재무비율", "Appendix A. Detailed Financial Ratios")
        
        ratios = analysis_data.get('ratios', {})
        
        if ratios:
            # 상세 재무비율 테이블
            ratio_table = doc.add_table(rows=len(ratios) + 1, cols=3)
            ratio_table.style = 'Table Grid'
            
            # 헤더
            header_cells = ratio_table.rows[0].cells
            headers = ["재무비율", "수치", "설명"]
            for i, header in enumerate(headers):
                para = header_cells[i].paragraphs[0]
                para.clear()
                run = para.add_run(header)
                run.font.bold = True
                run.font.name = 'Malgun Gothic'
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 데이터
            for i, (ratio_name, ratio_value) in enumerate(ratios.items(), 1):
                row = ratio_table.rows[i]
                
                # 비율명
                name_para = row.cells[0].paragraphs[0]
                name_para.clear()
                name_para.add_run(ratio_name).font.name = 'Malgun Gothic'
                
                # 수치
                value_para = row.cells[1].paragraphs[0]
                value_para.clear()
                if isinstance(ratio_value, (int, float)):
                    if ratio_name in ['ROE', 'ROA', '영업이익률', '순이익률', '부채비율']:
                        formatted_value = f"{ratio_value:.2f}%"
                    else:
                        formatted_value = f"{ratio_value:.2f}"
                else:
                    formatted_value = str(ratio_value)
                value_para.add_run(formatted_value).font.name = 'Malgun Gothic'
                
                # 설명
                desc_para = row.cells[2].paragraphs[0]
                desc_para.clear()
                description = self._get_ratio_description(ratio_name, ratio_value)
                desc_para.add_run(description).font.name = 'Malgun Gothic'

    def _add_section_heading(self, doc: Document, korean_title: str, english_title: str = ""):
        """섹션 제목 추가"""
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(korean_title)
        heading_run.font.name = 'Malgun Gothic'
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = self.color_scheme["primary"]
        
        if english_title:
            doc.add_paragraph()
            eng_para = doc.add_paragraph()
            eng_run = eng_para.add_run(english_title)
            eng_run.font.name = 'Calibri'
            eng_run.font.size = Pt(12)
            eng_run.font.italic = True
            eng_run.font.color.rgb = self.color_scheme["secondary"]
        
        doc.add_paragraph()

    def _add_subsection_heading(self, doc: Document, title: str):
        """하위 섹션 제목 추가"""
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(title)
        heading_run.font.name = 'Malgun Gothic'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = self.color_scheme["secondary"]
        doc.add_paragraph()

    def _format_currency(self, amount: int) -> str:
        """통화 포맷팅"""
        if abs(amount) >= 1000000000000:  # 1조 이상
            return f"{amount/1000000000000:.1f}조원"
        elif abs(amount) >= 100000000:  # 1억 이상
            return f"{amount/100000000:.0f}억원"
        elif abs(amount) >= 10000:  # 1만 이상
            return f"{amount/10000:.0f}만원"
        else:
            return f"{amount:,}원"

    def _get_ratio_description(self, ratio_name: str, ratio_value: Any) -> str:
        """재무비율 설명"""
        descriptions = {
            "ROE": "자기자본수익률 - 주주가 투자한 자본 대비 수익률",
            "ROA": "총자산수익률 - 전체 자산 대비 수익률",
            "영업이익률": "매출 대비 영업이익 비율 - 본업 수익성",
            "순이익률": "매출 대비 순이익 비율 - 최종 수익성",
            "부채비율": "자기자본 대비 부채 비율 - 재무 안정성",
            "자기자본비율": "총자산 대비 자기자본 비율 - 재무 건전성",
            "유동비율": "유동부채 대비 유동자산 비율 - 단기 지급능력",
            "총자산회전율": "총자산 대비 매출 효율성",
            "매출성장률": "전년 대비 매출 증감률",
            "순이익성장률": "전년 대비 순이익 증감률"
        }
        
        base_desc = descriptions.get(ratio_name, "재무 지표")
        
        # 평가 추가
        if isinstance(ratio_value, (int, float)):
            if ratio_name == "ROE":
                if ratio_value > 15:
                    base_desc += " (우수)"
                elif ratio_value > 8:
                    base_desc += " (양호)"
                else:
                    base_desc += " (개선필요)"
            elif ratio_name == "부채비율":
                if ratio_value < 100:
                    base_desc += " (안정)"
                elif ratio_value < 200:
                    base_desc += " (보통)"
                else:
                    base_desc += " (위험)"
        
        return base_desc

    def _add_image_to_document(self, doc: Document, image_path: str, caption: str = ""):
        """문서에 이미지 추가"""
        if os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6))
                if caption:
                    caption_para = doc.add_paragraph()
                    caption_run = caption_para.add_run(caption)
                    caption_run.font.name = 'Malgun Gothic'
                    caption_run.font.size = Pt(10)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except Exception as e:
                print(f"❌ 이미지 삽입 오류: {str(e)}")
        else:
            print(f"❌ 이미지 파일을 찾을 수 없습니다: {image_path}")

    def create_excel_report(self, analysis_data: Dict, company_name: str) -> str:
        """Excel 형태의 상세 분석 보고서"""
        try:
            import pandas as pd
            
            # Excel writer 생성
            filename = f"{company_name}_상세분석_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join(self.save_directory, filename)
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                
                # 1. 요약 시트
                summary_data = self._prepare_summary_data(analysis_data, company_name)
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='요약', index=False)
                
                # 2. 재무데이터 시트
                financial_data = analysis_data.get('financial_data', {})
                if financial_data:
                    fin_df = pd.DataFrame(list(financial_data.items()), 
                                        columns=['항목', '금액'])
                    fin_df.to_excel(writer, sheet_name='재무데이터', index=False)
                
                # 3. 재무비율 시트
                ratios = analysis_data.get('ratios', {})
                if ratios:
                    ratio_df = pd.DataFrame(list(ratios.items()), 
                                          columns=['비율명', '수치'])
                    ratio_df.to_excel(writer, sheet_name='재무비율', index=False)
                
                # 4. 부정위험 시트
                fraud_ratios = analysis_data.get('fraud_ratios', {})
                if fraud_ratios:
                    fraud_df = pd.DataFrame(list(fraud_ratios.items()), 
                                          columns=['위험지표', '값'])
                    fraud_df.to_excel(writer, sheet_name='부정위험', index=False)
            
            print(f"✅ Excel 보고서 생성: {filepath}")
            return filepath
            
        except ImportError:
            print("❌ pandas가 설치되지 않아 Excel 보고서를 생성할 수 없습니다.")
            return ""
        except Exception as e:
            print(f"❌ Excel 보고서 생성 오류: {str(e)}")
            return ""

    def _prepare_summary_data(self, analysis_data: Dict, company_name: str) -> List[Dict]:
        """요약 데이터 준비"""
        
        ratios = analysis_data.get('ratios', {})
        financial_data = analysis_data.get('financial_data', {})
        fraud_ratios = analysis_data.get('fraud_ratios', {})
        
        summary_data = [
            {"구분": "회사명", "내용": company_name},
            {"구분": "분석일", "내용": datetime.now().strftime("%Y-%m-%d")},
            {"구분": "매출액", "내용": self._format_currency(financial_data.get('revenue', 0))},
            {"구분": "순이익", "내용": self._format_currency(financial_data.get('net_income', 0))},
            {"구분": "ROE", "내용": f"{ratios.get('ROE', 0):.2f}%"},
            {"구분": "부채비율", "내용": f"{ratios.get('부채비율', 0):.2f}%"},
            {"구분": "부정위험점수", "내용": f"{fraud_ratios.get('종합_부정위험점수', 0):.0f}점"}
        ]
        
        return summary_data

    def create_pdf_report(self, analysis_data: Dict, company_name: str) -> str:
        """PDF 보고서 생성 (DOCX를 PDF로 변환)"""
        try:
            # 먼저 DOCX 생성
            docx_path = self.generate_comprehensive_report(analysis_data, company_name)
            
            # PDF 변환 시도 (python-docx2pdf 사용)
            try:
                from docx2pdf import convert
                pdf_path = docx_path.replace('.docx', '.pdf')
                convert(docx_path, pdf_path)
                print(f"✅ PDF 보고서 생성: {pdf_path}")
                return pdf_path
            except ImportError:
                print("⚠️ docx2pdf가 설치되지 않아 PDF 변환을 건너뜁니다.")
                print(f"✅ DOCX 보고서 생성: {docx_path}")
                return docx_path
            
        except Exception as e:
            print(f"❌ PDF 보고서 생성 오류: {str(e)}")
            return ""

    def generate_all_reports(self, analysis_data: Dict, company_name: str) -> Dict[str, str]:
        """모든 형태의 보고서 생성"""
        
        print(f"📋 {company_name} 전체 보고서 생성 시작...")
        
        reports = {}
        
        try:
            # 1. DOCX 보고서
            print("📄 DOCX 보고서 생성 중...")
            docx_path = self.generate_comprehensive_report(analysis_data, company_name)
            if docx_path:
                reports['docx'] = docx_path
            
            # 2. Excel 보고서
            print("📊 Excel 보고서 생성 중...")
            excel_path = self.create_excel_report(analysis_data, company_name)
            if excel_path:
                reports['excel'] = excel_path
            
            # 3. PDF 보고서 시도
            print("📑 PDF 보고서 생성 시도 중...")
            pdf_path = self.create_pdf_report(analysis_data, company_name)
            if pdf_path and pdf_path.endswith('.pdf'):
                reports['pdf'] = pdf_path
            
            print(f"✅ 총 {len(reports)}개 보고서 생성 완료")
            return reports
            
        except Exception as e:
            print(f"❌ 보고서 생성 중 오류: {str(e)}")
            return reports

    def open_report(self, filepath: str) -> bool:
        """생성된 보고서 열기"""
        try:
            import os
            import platform
            
            if os.path.exists(filepath):
                if platform.system() == 'Windows':
                    os.startfile(filepath)
                elif platform.system() == 'Darwin':  # macOS
                    os.system(f'open "{filepath}"')
                else:  # Linux
                    os.system(f'xdg-open "{filepath}"')
                return True
            else:
                print(f"❌ 파일을 찾을 수 없습니다: {filepath}")
                return False
                
        except Exception as e:
            print(f"❌ 파일 열기 오류: {str(e)}")
            return False

    def get_report_summary(self, reports: Dict[str, str]) -> str:
        """보고서 생성 결과 요약"""
        
        if not reports:
            return "❌ 생성된 보고서가 없습니다."
        
        summary = f"✅ {len(reports)}개 보고서 생성 완료:\n\n"
        
        for report_type, filepath in reports.items():
            if os.path.exists(filepath):
                file_size = os.path.getsize(filepath)
                size_mb = file_size / (1024 * 1024)
                
                summary += f"📋 {report_type.upper()} 보고서:\n"
                summary += f"   파일: {os.path.basename(filepath)}\n"
                summary += f"   크기: {size_mb:.2f} MB\n"
                summary += f"   경로: {filepath}\n\n"
        
        summary += "💡 보고서 파일들이 analysis_results/documents 폴더에 저장되었습니다."
        
        return summary